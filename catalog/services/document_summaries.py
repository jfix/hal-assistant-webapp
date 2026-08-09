from __future__ import annotations

import json
import logging
import os
import re
import time
from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader

MAX_UPLOAD_BYTES = 20 * 1024 * 1024
MAX_DOCUMENT_CHARACTERS = 120_000
MAX_PDF_PAGES = 300
MAX_DOCX_ENTRIES = 2_000
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 200
ALLOWED_EXTENSIONS = {".pdf", ".docx"}
SUMMARY_GENERATOR_VERSION = "humanities-bilingual-v2"
OPENAI_RETRY_DELAYS_SECONDS = (1.0, 3.0)

logger = logging.getLogger(__name__)


class DocumentSummaryError(Exception):
    """A safe error that can be shown to the person using the form."""


@dataclass(frozen=True)
class BilingualSummary:
    abstract_en: str
    abstract_fr: str
    keywords_en: list[str]
    keywords_fr: list[str]
    suggested_title: str = ""
    suggested_authors: list[str] | None = None
    suggested_publication_year: int | None = None
    suggested_publication_type: str = ""
    suggested_doi: str = ""


def document_sha256(upload) -> str:
    """Hash the original bytes and rewind the upload for subsequent extraction."""
    digest = sha256()
    for chunk in upload.chunks():
        digest.update(chunk)
    upload.seek(0)
    return digest.hexdigest()


def summary_model() -> str:
    return os.getenv("OPENAI_MODEL", "gpt-5.6-terra")


_TITLE_BOILERPLATE = re.compile(
    r"^(abstract|résumé|summary|mots[- ]clés|keywords?|key[- ]words?|copyright|"
    r"doi\b|https?://|www\.|issn\b|isbn\b)",
    re.IGNORECASE,
)
_FILELIKE_TITLE = re.compile(r"\.(indd?|docx?|pdf|odt)$", re.IGNORECASE)


def _clean_title(value: str) -> str:
    return " ".join(value.split()).strip("-–—:; ")


def _credible_title(value: str) -> bool:
    words = value.split()
    return bool(
        8 <= len(value) <= 300
        and 2 <= len(words) <= 40
        and not value.isdigit()
        and not _TITLE_BOILERPLATE.search(value)
        and not _FILELIKE_TITLE.search(value)
    )


def _title_score(value: str, position: int) -> int:
    if not _credible_title(value):
        return -10_000
    words = value.split()
    letters = [character for character in value if character.isalpha()]
    uppercase_ratio = (
        sum(character.isupper() for character in letters) / len(letters) if letters else 0
    )
    score = max(0, 40 - position)
    if 4 <= len(words) <= 18:
        score += 35
    if 20 <= len(value) <= 180:
        score += 25
    if uppercase_ratio >= 0.8 and len(words) >= 3:
        score += 60
    if value.endswith(("?", "!")) or ":" in value:
        score += 8
    if value.endswith(".") and len(words) > 12:
        score -= 25
    probable_person_name = (
        2 <= len(words) <= 4
        and uppercase_ratio < 0.8
        and all(word[:1].isupper() for word in words if word[:1].isalpha())
        and not any(mark in value for mark in (":", "?", "!", ","))
    )
    if probable_person_name:
        score -= 65
    return score


def infer_document_title(
    text: str,
    filename: str,
    preferred_titles: tuple[str, ...] = (),
) -> str:
    """Infer a title from credible metadata and scored opening-page lines."""
    for preferred in preferred_titles:
        candidate = _clean_title(preferred)
        if _credible_title(candidate):
            return candidate[:300]

    opening_lines = [_clean_title(line) for line in text.splitlines()[:40]]
    candidates = [
        (line, _title_score(line, position))
        for position, line in enumerate(opening_lines)
        if line
    ]
    if candidates:
        candidate, score = max(candidates, key=lambda item: item[1])
        if score > 0:
            return candidate[:300]
    fallback = Path(filename).stem.replace("_", " ").replace("-", " ")
    return " ".join(fallback.split())[:300] or "Document sans titre"


def extract_document_title(upload, text: str) -> str:
    """Read credible embedded/style title hints, then apply textual scoring."""
    extension = Path(upload.name).suffix.lower()
    upload.seek(0)
    contents = upload.read()
    upload.seek(0)
    preferred: list[str] = []
    try:
        if extension == ".docx":
            document = Document(BytesIO(contents))
            if document.core_properties.title:
                preferred.append(document.core_properties.title)
            preferred.extend(
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.text.strip()
                and paragraph.style.name.lower() in {"title", "titre"}
            )
        elif extension == ".pdf":
            metadata = PdfReader(BytesIO(contents)).metadata
            if metadata and metadata.title:
                preferred.append(metadata.title)
    except Exception:
        preferred = []
    return infer_document_title(text, upload.name, tuple(preferred))


def extract_document_text(upload) -> str:
    """Extract text from a PDF or DOCX upload without persisting the source file."""
    extension = Path(upload.name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise DocumentSummaryError("Formats acceptés : PDF (.pdf) et Word (.docx).")
    if upload.size > MAX_UPLOAD_BYTES:
        raise DocumentSummaryError("Le document dépasse la limite de 20 Mo.")

    contents = upload.read()
    try:
        if extension == ".pdf":
            if not contents.startswith(b"%PDF-"):
                raise DocumentSummaryError(
                    "Le contenu du fichier ne correspond pas à un document PDF."
                )
            reader = PdfReader(BytesIO(contents))
            if reader.is_encrypted:
                raise DocumentSummaryError(
                    "Les PDF protégés par mot de passe ne sont pas acceptés."
                )
            if len(reader.pages) > MAX_PDF_PAGES:
                raise DocumentSummaryError(f"Le PDF dépasse la limite de {MAX_PDF_PAGES} pages.")
            extracted_pages = []
            extracted_length = 0
            for page in reader.pages:
                page_text = page.extract_text() or ""
                extracted_pages.append(page_text)
                extracted_length += len(page_text)
                if extracted_length >= MAX_DOCUMENT_CHARACTERS:
                    break
            text = "\n\n".join(extracted_pages)
        else:
            _validate_docx_archive(contents)
            document = Document(BytesIO(contents))
            paragraphs = []
            extracted_length = 0
            for paragraph in document.paragraphs:
                paragraphs.append(paragraph.text)
                extracted_length += len(paragraph.text)
                if extracted_length >= MAX_DOCUMENT_CHARACTERS:
                    break
            text = "\n".join(paragraphs)
    except DocumentSummaryError:
        raise
    except Exception as exc:
        raise DocumentSummaryError(
            "Impossible de lire ce document. Vérifiez qu’il n’est pas chiffré ou endommagé."
        ) from exc

    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(text) < 200:
        raise DocumentSummaryError(
            "Le document ne contient pas assez de texte extractible. Les PDF numérisés "
            "nécessitent une étape OCR."
        )
    return text[:MAX_DOCUMENT_CHARACTERS]


def _validate_docx_archive(contents: bytes) -> None:
    if not contents.startswith(b"PK"):
        raise DocumentSummaryError(
            "Le contenu du fichier ne correspond pas à un document Word DOCX."
        )
    try:
        with ZipFile(BytesIO(contents)) as archive:
            entries = archive.infolist()
            names = {entry.filename for entry in entries}
            if {"[Content_Types].xml", "word/document.xml"} - names:
                raise DocumentSummaryError(
                    "Le fichier DOCX ne contient pas une structure Word valide."
                )
            if len(entries) > MAX_DOCX_ENTRIES:
                raise DocumentSummaryError("Le fichier DOCX contient trop d’éléments internes.")
            total_size = sum(entry.file_size for entry in entries)
            total_compressed = sum(entry.compress_size for entry in entries)
            if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise DocumentSummaryError("Le contenu décompressé du DOCX est trop volumineux.")
            if total_compressed and total_size / total_compressed > MAX_DOCX_COMPRESSION_RATIO:
                raise DocumentSummaryError("Le taux de compression du DOCX est jugé dangereux.")
            if archive.testzip() is not None:
                raise DocumentSummaryError("Le fichier DOCX est endommagé.")
    except BadZipFile as exc:
        raise DocumentSummaryError("Le fichier DOCX est endommagé.") from exc


def _response_text(payload: dict) -> str:
    if isinstance(payload.get("output_text"), str):
        return payload["output_text"]
    for item in payload.get("output", []):
        for content in item.get("content", []):
            if content.get("type") == "output_text" and isinstance(content.get("text"), str):
                return content["text"]
    raise DocumentSummaryError("Le service d’IA a renvoyé une réponse vide.")


def _retry_delay(exc: HTTPError | None, fallback: float) -> float:
    if exc is None:
        return fallback
    retry_after = exc.headers.get("Retry-After") if exc.headers else None
    try:
        return min(10.0, max(0.0, float(retry_after)))
    except (TypeError, ValueError):
        return fallback


def _openai_response(request: Request) -> dict:
    """Call OpenAI, retrying only failures that are expected to be transient."""
    for attempt in range(len(OPENAI_RETRY_DELAYS_SECONDS) + 1):
        try:
            with urlopen(request, timeout=90) as response:  # noqa: S310 - fixed HTTPS default
                return json.load(response)
        except HTTPError as exc:
            retryable = exc.code == 429 or exc.code >= 500
            if not retryable or attempt == len(OPENAI_RETRY_DELAYS_SECONDS):
                if exc.code == 401:
                    message = "La configuration du service d’IA doit être vérifiée."
                elif exc.code == 429:
                    message = "Le quota du service d’IA est temporairement indisponible ou épuisé."
                elif 400 <= exc.code < 500:
                    message = "Le service d’IA n’a pas accepté cette demande."
                else:
                    message = "Le service d’IA est momentanément indisponible."
                raise DocumentSummaryError(message) from exc
            delay = _retry_delay(exc, OPENAI_RETRY_DELAYS_SECONDS[attempt])
            logger.warning(
                "openai_transient_failure attempt=%s status=%s retry_in=%.1f",
                attempt + 1,
                exc.code,
                delay,
            )
            time.sleep(delay)
        except (URLError, TimeoutError) as exc:
            if attempt == len(OPENAI_RETRY_DELAYS_SECONDS):
                raise DocumentSummaryError(
                    "Le service d’IA est momentanément inaccessible."
                ) from exc
            delay = OPENAI_RETRY_DELAYS_SECONDS[attempt]
            logger.warning(
                "openai_transient_failure attempt=%s status=network retry_in=%.1f",
                attempt + 1,
                delay,
            )
            time.sleep(delay)
    raise AssertionError("unreachable")


def generate_bilingual_summary(text: str) -> BilingualSummary:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise DocumentSummaryError(
            "La génération IA n’est pas configurée. Définissez OPENAI_API_KEY puis relancez "
            "le serveur."
        )

    schema = {
        "type": "object",
        "properties": {
            "abstract_en": {"type": "string"},
            "abstract_fr": {"type": "string"},
            "keywords_en": {"type": "array", "items": {"type": "string"}, "maxItems": 10},
            "keywords_fr": {"type": "array", "items": {"type": "string"}, "maxItems": 10},
            "suggested_title": {"type": "string"},
            "suggested_authors": {"type": "array", "items": {"type": "string"}},
            "suggested_publication_year": {"type": ["integer", "null"]},
            "suggested_publication_type": {"type": "string"},
            "suggested_doi": {"type": "string"},
        },
        "required": [
            "abstract_en",
            "abstract_fr",
            "keywords_en",
            "keywords_fr",
            "suggested_title",
            "suggested_authors",
            "suggested_publication_year",
            "suggested_publication_type",
            "suggested_doi",
        ],
        "additionalProperties": False,
    }
    request_body = {
        "model": summary_model(),
        "reasoning": {"effort": "low"},
        "instructions": (
            "You are an expert academic editor in the humanities. Base every claim only on the "
            "supplied document. Write two equivalent, publication-ready abstracts, one in English "
            "and one in French, of 150–220 words each. Preserve nuance, method, corpus, central "
            "argument, and contribution. Do not invent missing facts. Extract at most ten concise "
            "concepts or keywords in each language; the two lists must be aligned translations in "
            "the same order. Also extract bibliographic suggestions: the document title, named "
            "authors, four-digit publication year, DOI, and a conservative publication type "
            "(article, book_chapter, conference_paper, book, edited_book, other, or empty). "
            "Use empty values when the document does not establish a fact. Return only the "
            "requested structured data."
        ),
        "input": text,
        "text": {
            "format": {
                "type": "json_schema",
                "name": "bilingual_article_summary",
                "strict": True,
                "schema": schema,
            }
        },
        "store": False,
    }
    request = Request(
        os.getenv("OPENAI_RESPONSES_URL", "https://api.openai.com/v1/responses"),
        data=json.dumps(request_body).encode(),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    response_payload = _openai_response(request)

    try:
        result = json.loads(_response_text(response_payload))
        return BilingualSummary(
            abstract_en=result["abstract_en"].strip(),
            abstract_fr=result["abstract_fr"].strip(),
            keywords_en=[item.strip() for item in result["keywords_en"][:10]],
            keywords_fr=[item.strip() for item in result["keywords_fr"][:10]],
            suggested_title=result["suggested_title"].strip(),
            suggested_authors=[
                item.strip() for item in result["suggested_authors"] if item.strip()
            ],
            suggested_publication_year=result["suggested_publication_year"],
            suggested_publication_type=result["suggested_publication_type"].strip(),
            suggested_doi=result["suggested_doi"].strip(),
        )
    except (KeyError, TypeError, json.JSONDecodeError) as exc:
        raise DocumentSummaryError("La réponse de l’IA n’a pas le format attendu.") from exc
