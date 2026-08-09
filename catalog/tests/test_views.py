from __future__ import annotations

import re
from io import BytesIO
from unittest.mock import patch

import pytest
from django.contrib.auth import get_user_model
from django.contrib.auth.models import Permission
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from django.utils import timezone
from docx import Document
from openpyxl import load_workbook

from catalog.models import (
    DocumentPublicationLink,
    DocumentSummaryCache,
    HALOperation,
    Publication,
    SourceImport,
    SourceRecord,
)
from catalog.services.document_summaries import BilingualSummary

pytestmark = pytest.mark.django_db


@pytest.fixture
def user():
    return get_user_model().objects.create_user(
        username="reviewer",
        password="local-test-password",
    )


@pytest.fixture
def publications() -> tuple[Publication, Publication]:
    first = Publication.objects.create(
        publication_key="pub-0001",
        publication_type="journal_article",
        title="Portable metadata",
        publication_year=2025,
        authors=["Ada Lovelace"],
        readiness_state=Publication.ReadinessState.HAL_READY,
    )
    second = Publication.objects.create(
        publication_key="pub-0002",
        publication_type="conference_paper",
        title="Cloud deployment",
        publication_year=2024,
        authors=["Grace Hopper"],
        readiness_state=Publication.ReadinessState.NEEDS_ENRICHMENT,
        missing_required_fields=["conference_city"],
    )
    return first, second


def test_health_is_public_and_checks_database(client) -> None:
    response = client.get(reverse("health"))

    assert response.status_code == 200
    assert response.json() == {"status": "ok", "database": "reachable"}


def test_home_is_a_dashboard_with_stats_and_recent_drafts(client, user) -> None:
    client.force_login(user)
    draft = Publication.objects.create(
        publication_key="pub-dashboard-draft",
        publication_type="journal_article",
        title="Notice à terminer",
        missing_required_fields=["journal_title"],
    )
    Publication.objects.create(
        publication_key="pub-dashboard-hal",
        publication_type="journal_article",
        title="Notice déjà publiée",
        hal_id="hal-01234567",
    )

    response = client.get(reverse("home"))
    content = response.content.decode()

    assert response.status_code == 200
    assert "Tableau de bord" in content
    assert "Ajouter une publication" in content
    assert "data-local-greeting" in content
    assert "Notice à terminer" in content
    assert "Notice déjà publiée" not in content
    assert reverse("publication-detail", args=[draft.id]) in content
    assert f'{reverse("publication-list")}?workflow=draft' in content


def test_publication_list_can_filter_by_deposit_state(client, user) -> None:
    client.force_login(user)
    Publication.objects.create(
        publication_key="pub-filter-draft",
        publication_type="book",
        title="Brouillon visible",
    )
    Publication.objects.create(
        publication_key="pub-filter-hal",
        publication_type="book",
        title="Publication masquée",
        hal_id="hal-07654321",
    )

    response = client.get(reverse("publication-list"), {"workflow": "draft"})
    content = response.content.decode()

    assert response.status_code == 200
    assert "Brouillon visible" in content
    assert "Publication masquée" not in content
    assert '<option value="draft" selected>Brouillons</option>' in content


@pytest.mark.parametrize(
    ("content_missing", "visible_title", "hidden_title"),
    [
        ("abstracts", "Résumé français manquant", "Notice bilingue complète"),
        ("abstract_fr", "Résumé français manquant", "Mots-clés anglais manquants"),
        ("keywords", "Mots-clés anglais manquants", "Notice bilingue complète"),
        ("keywords_en", "Mots-clés anglais manquants", "Résumé français manquant"),
        ("bilingual_content", "Résumé français manquant", "Notice bilingue complète"),
    ],
)
def test_publication_list_can_filter_missing_bilingual_content(
    client,
    user,
    content_missing: str,
    visible_title: str,
    hidden_title: str,
) -> None:
    client.force_login(user)
    common = {
        "publication_type": "journal_article",
        "abstract_en": "English abstract",
        "keywords_fr": ["mémoire"],
    }
    Publication.objects.create(
        publication_key="pub-content-complete",
        title="Notice bilingue complète",
        abstract_fr="Résumé français",
        keywords_en=["memory"],
        **common,
    )
    Publication.objects.create(
        publication_key="pub-missing-abstract-fr",
        title="Résumé français manquant",
        abstract_fr="",
        keywords_en=["memory"],
        **common,
    )
    Publication.objects.create(
        publication_key="pub-missing-keywords-en",
        title="Mots-clés anglais manquants",
        abstract_fr="Résumé français",
        keywords_en=[],
        **common,
    )

    response = client.get(
        reverse("publication-list"),
        {"content_missing": content_missing},
    )
    content = response.content.decode()

    assert response.status_code == 200
    assert visible_title in content
    assert hidden_title not in content
    assert f'<option value="{content_missing}" selected>' in content


@pytest.mark.parametrize("name", ["publication-list", "home"])
def test_catalog_routes_require_authentication(client, name: str) -> None:
    response = client.get(reverse(name))

    assert response.status_code == 302
    assert response.url.startswith(reverse("login"))


def test_document_summary_requires_authentication(client) -> None:
    response = client.get(reverse("document-summary"))

    assert response.status_code == 302
    assert response.url.startswith(reverse("login"))


def test_admin_navigation_is_visible_only_to_staff(client, user) -> None:
    client.force_login(user)
    regular_response = client.get(reverse("document-summary"))

    user.is_staff = True
    user.save(update_fields=["is_staff"])
    staff_response = client.get(reverse("document-summary"))

    admin_url = reverse("admin:index")
    assert f'href="{admin_url}"' not in regular_response.content.decode()
    assert f'href="{admin_url}"' in staff_response.content.decode()
    assert "Administration" in staff_response.content.decode()


def test_application_responses_include_restrictive_security_headers(client, user) -> None:
    client.force_login(user)

    response = client.get(reverse("document-summary"))

    csp = response["Content-Security-Policy"]
    assert "default-src 'self'" in csp
    assert "script-src 'self';" in csp
    assert "unsafe-inline" not in csp.split("style-src")[0]
    assert response["Referrer-Policy"] == "same-origin"
    assert "camera=()" in response["Permissions-Policy"]
    assert response["Cross-Origin-Opener-Policy"] == "same-origin"


def test_admin_csp_allows_legacy_inline_admin_scripts(client) -> None:
    admin_user = get_user_model().objects.create_superuser(
        username="header-admin",
        password="local-test-password",
    )
    client.force_login(admin_user)

    response = client.get(reverse("admin:index"))

    assert "script-src 'self' 'unsafe-inline'" in response["Content-Security-Policy"]


def test_document_summary_page_accepts_docx(client, user) -> None:
    document = Document()
    document.add_paragraph("A humanities argument about archives and cultural memory. " * 12)
    file_contents = BytesIO()
    document.save(file_contents)
    upload = SimpleUploadedFile(
        "article.docx",
        file_contents.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    generated = BilingualSummary(
        abstract_en="English abstract.",
        abstract_fr="Résumé français.",
        keywords_en=["archives", "memory"],
        keywords_fr=["archives", "mémoire"],
    )
    client.force_login(user)

    with patch(
        "catalog.services.publication_documents.generate_bilingual_summary", return_value=generated
    ) as generate:
        response = client.post(reverse("document-summary"), {"document": upload})

    assert response.status_code == 200
    content = response.content.decode()
    assert "Résumé français." in content
    assert "English abstract." in content
    assert "archives, mémoire" in content
    assert "cultural memory" in generate.call_args.args[0]
    assert DocumentSummaryCache.objects.count() == 1
    cache_entry = DocumentSummaryCache.objects.get()
    assert cache_entry.source_filename == "article.docx"
    assert cache_entry.document_title


def test_document_summary_reuses_cached_result(client, user) -> None:
    document = Document()
    document.add_paragraph("An analysis of archives, memory, and material history. " * 12)
    file_contents = BytesIO()
    document.save(file_contents)
    file_bytes = file_contents.getvalue()
    generated = BilingualSummary(
        abstract_en="Cached English abstract.",
        abstract_fr="Résumé français en cache.",
        keywords_en=["archives"],
        keywords_fr=["archives"],
    )
    client.force_login(user)

    with patch(
        "catalog.services.publication_documents.generate_bilingual_summary",
        return_value=generated,
    ) as generate:
        first = client.post(
            reverse("document-summary"),
            {"document": SimpleUploadedFile("first-name.docx", file_bytes)},
        )
        second = client.post(
            reverse("document-summary"),
            {"document": SimpleUploadedFile("renamed.docx", file_bytes)},
        )

    assert first.status_code == second.status_code == 200
    assert generate.call_count == 1
    assert "aucun appel API" in second.content.decode()
    assert DocumentSummaryCache.objects.count() == 1


def test_document_summary_rejects_unsupported_upload(client, user) -> None:
    client.force_login(user)
    upload = SimpleUploadedFile("article.txt", b"plain text" * 100)

    response = client.post(reverse("document-summary"), {"document": upload})

    assert response.status_code == 200
    assert "Formats acceptés" in response.content.decode()


def test_authenticated_user_can_filter_publications(client, user, publications) -> None:
    client.force_login(user)

    response = client.get(
        reverse("publication-list"),
        {"q": "Portable", "readiness": Publication.ReadinessState.HAL_READY},
    )

    assert response.status_code == 200
    assert publications[0].title in response.content.decode()
    assert publications[1].title not in response.content.decode()


def test_authenticated_user_can_view_publication_detail(
    client,
    user,
    publications,
) -> None:
    client.force_login(user)

    response = client.get(
        reverse("publication-detail", args=[publications[1].id]),
    )

    assert response.status_code == 200
    content = response.content.decode()
    assert "Cloud deployment" in content
    assert "conference_city" in content
    assert "Résumés et mots-clés" in content
    assert 'name="field" value="abstract_fr"' not in content
    header = content.split('<div class="detail-grid">', 1)[0]
    assert publications[1].publication_key not in header
    assert "<span>Notice</span>" in header


def test_detail_statuses_separate_hal_data_and_synchronization(client, user) -> None:
    draft = Publication.objects.create(
        publication_key="status-draft",
        publication_type="journal_article",
        title="Ready draft",
        readiness_state=Publication.ReadinessState.HAL_READY,
        review_state=Publication.ReviewState.APPROVED,
    )
    published = Publication.objects.create(
        publication_key="status-published",
        publication_type="journal_article",
        title="Published and modified",
        hal_id="hal-123456",
        readiness_state=Publication.ReadinessState.HAL_READY,
        version=3,
        hal_synced_version=2,
        review_state=Publication.ReviewState.APPROVED,
    )
    client.force_login(user)

    draft_content = client.get(
        reverse("publication-detail", args=[draft.id])
    ).content.decode()
    published_content = client.get(
        reverse("publication-detail", args=[published.id])
    ).content.decode()

    assert 'class="status-dimension">HAL</span>' in draft_content
    assert 'class="status-value">Brouillon</span>' in draft_content
    assert 'class="status-value">Prêt pour HAL</span>' in draft_content
    assert 'class="status-value">Jamais synchronisé</span>' in draft_content
    assert "Approuvé" not in draft_content
    assert 'class="status-value">Publié sur HAL</span>' in published_content
    assert 'class="status-value">Prêt pour mise à jour HAL</span>' in published_content
    assert 'class="status-value">Modifié</span>' in published_content
    assert "Approuvé" not in published_content


def test_published_unmodified_record_is_shown_as_synchronized(client, user) -> None:
    publication = Publication.objects.create(
        publication_key="status-synced",
        publication_type="journal_article",
        title="Published and synchronized",
        hal_id="hal-654321",
        version=4,
        hal_synced_version=4,
        readiness_state=Publication.ReadinessState.HAL_READY,
    )
    client.force_login(user)

    content = client.get(
        reverse("publication-detail", args=[publication.id])
    ).content.decode()

    assert 'class="status-value">Publié sur HAL</span>' in content
    assert 'class="status-value">Minimum HAL atteint</span>' in content
    assert 'class="status-value">À jour</span>' in content
    assert "Prêt pour mise à jour HAL" not in content
    assert 'class="status-value">Modifié</span>' not in content


def test_publication_list_uses_compact_three_column_statuses(client, user) -> None:
    Publication.objects.create(
        publication_key="status-list",
        publication_type="journal_article",
        title="Compact list statuses",
        hal_id="hal-765432",
        readiness_state=Publication.ReadinessState.HAL_READY,
        version=2,
        hal_synced_version=2,
    )
    client.force_login(user)

    content = client.get(reverse("publication-list")).content.decode()

    assert "status-summary-compact" in content
    assert 'class="status-dimension">Sync.</span>' in content
    assert 'class="status-value">Publié</span>' in content
    assert 'class="status-value">Minimum HAL atteint</span>' in content
    assert "Prêt à mettre à jour" not in content


def test_publication_list_displays_hal_type_codes_with_french_tooltips(
    client,
    user,
) -> None:
    client.force_login(user)
    Publication.objects.create(
        publication_key="pub-hal-type-explicit",
        publication_type="edited_book",
        hal_document_type="DOUV",
        title="Direction scientifique",
    )
    Publication.objects.create(
        publication_key="pub-hal-type-fallback",
        publication_type="journal_article",
        title="Article sans type matérialisé",
    )

    content = client.get(reverse("publication-list")).content.decode()

    assert '<abbr class="hal-type" title="Direction d’ouvrage">DOUV</abbr>' in content
    assert '<abbr class="hal-type" title="Article dans une revue">ART</abbr>' in content
    assert ">edited_book</td>" not in content


def test_publication_list_type_filter_uses_hal_taxonomy(client, user) -> None:
    client.force_login(user)
    Publication.objects.create(
        publication_key="pub-type-douv",
        publication_type="edited_book",
        hal_document_type="DOUV",
        title="Direction retenue",
    )
    Publication.objects.create(
        publication_key="pub-type-ouv",
        publication_type="book",
        hal_document_type="OUV",
        title="Ouvrage exclu",
    )
    Publication.objects.create(
        publication_key="pub-type-art-fallback",
        publication_type="journal_article",
        title="Article au type inféré",
    )

    douv_content = client.get(
        reverse("publication-list"), {"type": "DOUV"}
    ).content.decode()
    art_content = client.get(
        reverse("publication-list"), {"type": "ART"}
    ).content.decode()

    assert "Direction retenue" in douv_content
    assert "Ouvrage exclu" not in douv_content
    assert '<option value="DOUV" selected>DOUV — Direction d’ouvrage</option>' in douv_content
    assert "Article au type inféré" in art_content
    assert "Direction retenue" not in art_content


def test_publication_list_headers_sort_full_filtered_results_and_toggle_direction(
    client,
    user,
) -> None:
    client.force_login(user)
    Publication.objects.create(
        publication_key="pub-sort-zulu",
        publication_type="book",
        title="Zulu notice",
        publication_year=2020,
    )
    Publication.objects.create(
        publication_key="pub-sort-alpha",
        publication_type="book",
        title="Alpha notice",
        publication_year=2025,
    )

    ascending = client.get(
        reverse("publication-list"),
        {"workflow": "draft", "sort": "title", "direction": "asc"},
    ).content.decode()
    descending = client.get(
        reverse("publication-list"),
        {"workflow": "draft", "sort": "title", "direction": "desc"},
    ).content.decode()

    assert ascending.index("Alpha notice") < ascending.index("Zulu notice")
    assert descending.index("Zulu notice") < descending.index("Alpha notice")
    assert 'aria-sort="ascending"' in ascending
    assert "workflow=draft&amp;sort=title&amp;direction=desc" in ascending
    assert '>Publication <span aria-hidden="true">↑</span>' in ascending


def test_publication_list_exposes_sort_controls_for_every_visible_column(
    client,
    user,
) -> None:
    client.force_login(user)

    content = client.get(reverse("publication-list")).content.decode()

    for field in ("title", "year", "type", "state", "hal"):
        assert f"sort={field}&amp;direction=asc" in content


def test_publication_detail_reads_latest_hal_operation(client, user, publications) -> None:
    publication = publications[0]
    operation = HALOperation.objects.create(
        publication=publication,
        requested_by=user,
        publication_version=publication.version,
        state=HALOperation.State.PREPARED,
        duplicate_check={},
    )
    client.force_login(user)

    response = client.get(reverse("publication-detail", args=[publication.id]))

    assert response.status_code == 200
    assert reverse("hal-preprod-operation", args=[operation.id]) in response.content.decode()


def test_external_links_open_in_safe_new_tabs(client, user) -> None:
    publication = Publication.objects.create(
        publication_key="external-links",
        publication_type="journal_article",
        title="External link policy",
        doi="10.1234/example",
        isbn=["978-1-2345-6789-0"],
        source_url="https://journal.example/article",
        hal_id="hal-123456",
    )
    client.force_login(user)

    detail = client.get(reverse("publication-detail", args=[publication.id])).content.decode()
    listing = client.get(reverse("publication-list")).content.decode()

    external_anchors = re.findall(r'<a\s+[^>]*href="https?://[^>]+>', detail + listing)
    assert len(external_anchors) == 5
    for anchor in external_anchors:
        assert 'target="_blank"' in anchor
        assert 'rel="noopener noreferrer"' in anchor


@pytest.mark.parametrize(
    "action,label",
    [
        (DocumentPublicationLink.Action.LINKED, "Notice existante"),
        (DocumentPublicationLink.Action.CREATED, "Nouveau brouillon local"),
    ],
)
def test_publication_detail_shows_linked_document_analysis_before_sources(
    client, user, publications, action, label
) -> None:
    summary = DocumentSummaryCache.objects.create(
        owner=user,
        source_filename="article.pdf",
        document_title="Document analysé",
        document_sha256="c" * 64,
        model_name="test-model",
        generator_version="test-version",
        abstract_en="Associated English abstract.",
        abstract_fr="Résumé français associé.",
        keywords_en=["archive", "memory"],
        keywords_fr=["archives", "mémoire"],
    )
    DocumentPublicationLink.objects.create(
        summary=summary,
        publication=publications[0],
        actor=user,
        action=action,
    )
    client.force_login(user)

    response = client.get(reverse("publication-detail", args=[publications[0].id]))

    content = response.content.decode()
    assert response.status_code == 200
    assert "Analyses de documents associées" in content
    assert "article.pdf" in content
    assert "Résumé français associé." in content
    assert "Associated English abstract." in content
    assert "archives, mémoire" in content
    assert label in content
    assert "Ouvrir le résultat en cache" in content
    assert content.index("Analyses de documents associées") < content.index(
        "Sources d'origine"
    )


def test_publication_detail_does_not_link_to_another_users_private_cache(
    client, user, publications
) -> None:
    owner = get_user_model().objects.create_user(username="owner", password="password")
    summary = DocumentSummaryCache.objects.create(
        owner=owner,
        document_title="Private cache",
        document_sha256="d" * 64,
        model_name="test-model",
        generator_version="test-version",
        abstract_en="English",
        abstract_fr="Français",
    )
    DocumentPublicationLink.objects.create(
        summary=summary,
        publication=publications[0],
        actor=owner,
        action=DocumentPublicationLink.Action.LINKED,
    )
    client.force_login(user)

    response = client.get(reverse("publication-detail", args=[publications[0].id]))

    assert "Private cache" in response.content.decode()
    assert "Ouvrir le résultat en cache" not in response.content.decode()


def test_authenticated_user_can_open_another_users_associated_source_document(
    client, user, publications, settings, tmp_path
) -> None:
    settings.MEDIA_ROOT = tmp_path
    owner = get_user_model().objects.create_user(
        username="document-owner", password="password"
    )
    summary = DocumentSummaryCache.objects.create(
        owner=owner,
        source_filename="original-article.pdf",
        source_file=SimpleUploadedFile(
            "original-article.pdf",
            b"%PDF-1.4 safe test document",
            content_type="application/pdf",
        ),
        document_sha256="e" * 64,
        model_name="test-model",
        generator_version="test-version",
        abstract_en="English",
        abstract_fr="Français",
    )
    DocumentPublicationLink.objects.create(
        summary=summary,
        publication=publications[0],
        actor=owner,
        action=DocumentPublicationLink.Action.LINKED,
    )
    client.force_login(user)

    detail = client.get(reverse("publication-detail", args=[publications[0].id]))
    response = client.get(reverse("associated-document", args=[summary.id]))

    assert "Ouvrir le document original" in detail.content.decode()
    assert 'target="_blank"' in detail.content.decode()
    assert response.status_code == 200
    assert response["Content-Type"] == "application/pdf"
    assert response["Cache-Control"] == "private, no-store"
    assert b"".join(response.streaming_content) == b"%PDF-1.4 safe test document"


def test_unassociated_source_document_cannot_be_opened(
    client, user, settings, tmp_path
) -> None:
    settings.MEDIA_ROOT = tmp_path
    summary = DocumentSummaryCache.objects.create(
        owner=user,
        source_filename="unlinked.pdf",
        source_file=SimpleUploadedFile("unlinked.pdf", b"%PDF-1.4 unlinked"),
        document_sha256="f" * 64,
        model_name="test-model",
        generator_version="test-version",
        abstract_en="English",
        abstract_fr="Français",
    )
    client.force_login(user)

    response = client.get(reverse("associated-document", args=[summary.id]))

    assert response.status_code == 404


def test_associated_source_document_requires_authentication(client) -> None:
    response = client.get(
        reverse("associated-document", args=["00000000-0000-0000-0000-000000000000"])
    )

    assert response.status_code == 302
    assert response.url.startswith(reverse("login"))


def test_detail_page_has_no_unrendered_template_syntax(client, user, publications) -> None:
    """Guard against leaked template comments/tags (e.g. multi-line {# #})."""
    client.force_login(user)

    response = client.get(reverse("publication-detail", args=[publications[0].id]))

    content = response.content.decode()
    assert response.status_code == 200
    for leak in ("{#", "{%", "Renders one metadata field", "surrounding context"):
        assert leak not in content


def test_publication_meta_sections_are_collapsed_by_default(
    client, user, publications
) -> None:
    client.force_login(user)

    response = client.get(reverse("publication-detail", args=[publications[0].id]))

    content = response.content.decode()
    assert content.count('<details class="panel meta-panel">') == 3
    assert '<details class="panel meta-panel" open>' not in content
    assert "Sources d'origine" in content
    assert "Assertions de champs" in content
    assert "Historique des décisions" in content


def test_export_returns_a_filtered_xlsx_snapshot(client, user, publications) -> None:
    client.force_login(user)

    response = client.get(reverse("publication-export"))

    assert response.status_code == 200
    assert response["Content-Type"].startswith(
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    assert "attachment" in response["Content-Disposition"]
    sheet = load_workbook(BytesIO(response.content))["Publications"]
    header = [cell.value for cell in sheet[1]]
    assert header[0] == "publication_id"
    assert "title" in header
    keys = {row[0] for row in sheet.iter_rows(min_row=2, values_only=True)}
    assert {"pub-0001", "pub-0002"} <= keys


def test_export_respects_active_filters(client, user, publications) -> None:
    client.force_login(user)

    response = client.get(reverse("publication-export"), {"type": "COMM"})

    sheet = load_workbook(BytesIO(response.content))["Publications"]
    keys = [row[0] for row in sheet.iter_rows(min_row=2, values_only=True)]
    assert keys == ["pub-0002"]


def test_export_requires_authentication(client) -> None:
    response = client.get(reverse("publication-export"))

    assert response.status_code == 302
    assert response.url.startswith(reverse("login"))


@pytest.fixture
def publication_with_source() -> Publication:
    publication = Publication.objects.create(
        publication_key="pub-xml-1",
        publication_type="journal_article",
        hal_document_type="ART",
        title="A debuggable notice",
        publication_year=2024,
        authors=["Ada Lovelace"],
    )
    source_import = SourceImport.objects.create(
        source_type=SourceImport.SourceType.XLSX,
        source_name="review.xlsx",
        stored_file="snapshots/ab/deadbeef.xlsx",
        file_sha256="d" * 64,
        parser_version="hal-assistant/test",
        report_sha256="e" * 64,
        record_count=1,
        report={},
        retrieved_at=timezone.now(),
    )
    SourceRecord.objects.create(
        source_import=source_import,
        publication=publication,
        locator="Publications!row-2",
        original_citation="Lovelace, Ada. A debuggable notice. 2024.",
        raw_data={
            "publication_id": "pub-xml-1",
            "title": "A debuggable notice",
            "document_type": "ART",
            "year": 2024,
            "authors": "Ada Lovelace",
            "hal_domain": "shs.litt",
            "idhal": "florence-fix",
        },
        record_sha256="f" * 64,
    )
    return publication


def test_submission_xml_debug_view_renders_tei(
    client,
    user,
    publication_with_source,
) -> None:
    client.force_login(user)

    response = client.get(
        reverse("publication-xml", args=[publication_with_source.id]),
    )

    assert response.status_code == 200
    content = response.content.decode()
    # The notice is HTML-escaped for safe display on the page.
    assert "&lt;TEI" in content
    assert "Aperçu de débogage" in content


def test_submission_xml_raw_format_returns_xml(
    client,
    user,
    publication_with_source,
) -> None:
    client.force_login(user)

    response = client.get(
        reverse("publication-xml", args=[publication_with_source.id]),
        {"format": "raw"},
    )

    assert response.status_code == 200
    assert response["Content-Type"].startswith("application/xml")
    assert b"<TEI" in response.content


def test_submission_xml_uses_current_multilingual_abstracts_and_keyword_terms(
    client,
    user,
    publication_with_source,
) -> None:
    publication_with_source.abstract_fr = "Résumé actuel."
    publication_with_source.abstract_en = "Current abstract."
    publication_with_source.keywords_fr = ["archives", "mémoire, histoire"]
    publication_with_source.keywords_en = ["archives", "cultural memory"]
    publication_with_source.save()
    client.force_login(user)

    response = client.get(
        reverse("publication-xml", args=[publication_with_source.id]),
        {"format": "raw"},
    )

    xml = response.content.decode()
    assert '<keywords scheme="author">' in xml
    assert '<term xml:lang="fr">mémoire, histoire</term>' in xml
    assert '<term xml:lang="en">cultural memory</term>' in xml
    assert '<abstract xml:lang="fr">' in xml
    assert "Résumé actuel." in xml


def test_keyword_fields_render_as_manually_editable_pills(
    client,
    user,
    publications,
) -> None:
    publication = publications[0]
    publication.keywords_fr = ["archives", "mémoire, histoire"]
    publication.save()
    user.user_permissions.add(Permission.objects.get(codename="review_publication"))
    client.force_login(user)

    response = client.get(reverse("publication-detail", args=[publication.id]))

    content = response.content.decode()
    assert '<span class="keyword-pill">archives</span>' in content
    assert '<span class="keyword-pill">mémoire, histoire</span>' in content
    assert "data-keyword-input" in content
    assert "data-keyword-add" in content
    assert "data-keyword-save" in content
    assert "data-keyword-cancel" in content


def test_submission_xml_requires_authentication(client, publication_with_source) -> None:
    response = client.get(
        reverse("publication-xml", args=[publication_with_source.id]),
    )

    assert response.status_code == 302
    assert response.url.startswith(reverse("login"))
