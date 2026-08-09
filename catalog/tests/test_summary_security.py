from __future__ import annotations

import json
from datetime import timedelta
from io import BytesIO
from unittest.mock import patch
from urllib.error import HTTPError, URLError
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.management import call_command
from django.urls import reverse
from django.utils import timezone
from docx import Document
from pypdf import PdfWriter

from catalog.models import (
    ActiveDocumentSummaryGeneration,
    AuditEvent,
    DocumentSummaryCache,
    DocumentSummaryGenerationAttempt,
)
from catalog.services.document_summaries import (
    BilingualSummary,
    DocumentSummaryError,
    extract_document_text,
    generate_bilingual_summary,
    infer_document_title,
)
from catalog.services.summary_cache import purge_expired_summary_cache

pytestmark = pytest.mark.django_db


def _user(username: str):
    return get_user_model().objects.create_user(username=username, password="safe-test-password")


def _docx_bytes(marker: str) -> bytes:
    document = Document()
    document.add_paragraph((f"A humanities study of {marker}, archives, and memory. ") * 12)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _upload(contents: bytes, name: str = "article.docx") -> SimpleUploadedFile:
    return SimpleUploadedFile(name, contents)


def _summary(marker: str) -> BilingualSummary:
    return BilingualSummary(
        abstract_en=f"English {marker}.",
        abstract_fr=f"Français {marker}.",
        keywords_en=[marker],
        keywords_fr=[marker],
    )


def test_cache_is_scoped_to_authenticated_user(client) -> None:
    first_user = _user("first")
    second_user = _user("second")
    contents = _docx_bytes("shared document")

    with patch(
        "catalog.services.publication_documents.generate_bilingual_summary",
        side_effect=[_summary("first"), _summary("second")],
    ) as generate:
        client.force_login(first_user)
        client.post(reverse("document-summary"), {"document": _upload(contents)})
        client.force_login(second_user)
        response = client.post(reverse("document-summary"), {"document": _upload(contents)})

    assert generate.call_count == 2
    assert "Français second." in response.content.decode()
    assert set(DocumentSummaryCache.objects.values_list("owner__username", flat=True)) == {
        "first",
        "second",
    }


def test_document_title_is_inferred_locally_with_filename_fallback() -> None:
    assert infer_document_title("A Cultural History of Memory\nBody text", "article.pdf") == (
        "A Cultural History of Memory"
    )
    assert infer_document_title("", "my_archival-study.docx") == "my archival study"


def test_title_inference_skips_abstracts_keywords_and_front_matter() -> None:
    long_abstract = "This opening abstract sentence is intentionally long. " * 10
    text = "\n".join(
        [
            long_abstract,
            "Mots-clés : mélodrame ; théâtre ; traduction",
            "Key-words: melodrama; theatre; translation",
            "A Hidden Passenger in European Romantic Theatre?",
            "The body of the article begins here with a long explanatory sentence.",
        ]
    )

    assert infer_document_title(text, "article.docx") == (
        "A Hidden Passenger in European Romantic Theatre?"
    )


def test_title_inference_prefers_all_caps_heading_over_copyright_and_author() -> None:
    text = "\n".join(
        [
            "Copyright University Press",
            "177",
            "Florence Example",
            "TYRANNICIDE AND SUPERVISED FREEDOM",
            "The article begins with a substantial sentence about political liberty.",
        ]
    )

    assert infer_document_title(text, "chapter.pdf", ("chapter.indb",)) == (
        "TYRANNICIDE AND SUPERVISED FREEDOM"
    )


def test_cache_hit_does_not_consume_generation_quota(client, monkeypatch) -> None:
    monkeypatch.setenv("SUMMARY_USER_MINUTE_LIMIT", "1")
    user = _user("cached")
    contents = _docx_bytes("cache quota")
    client.force_login(user)

    with patch(
        "catalog.services.publication_documents.generate_bilingual_summary",
        return_value=_summary("cached"),
    ) as generate:
        first = client.post(reverse("document-summary"), {"document": _upload(contents)})
        second = client.post(reverse("document-summary"), {"document": _upload(contents)})

    assert first.status_code == second.status_code == 200
    assert generate.call_count == 1
    assert DocumentSummaryGenerationAttempt.objects.filter(owner=user).count() == 1
    assert "aucun appel API" in second.content.decode()


def test_cache_hit_refreshes_inferred_title_without_api_call(client) -> None:
    user = _user("title-refresh")
    document = Document()
    document.add_paragraph("The Correct Document Title", style="Title")
    document.add_paragraph("A humanities discussion of archives and memory. " * 12)
    output = BytesIO()
    document.save(output)
    contents = output.getvalue()
    client.force_login(user)

    with patch(
        "catalog.services.publication_documents.generate_bilingual_summary",
        return_value=_summary("refresh"),
    ) as generate:
        client.post(reverse("document-summary"), {"document": _upload(contents)})
        entry = DocumentSummaryCache.objects.get(owner=user)
        DocumentSummaryCache.objects.filter(id=entry.id).update(document_title="Wrong title")
        response = client.post(
            reverse("document-summary"),
            {"document": _upload(contents, "renamed.docx")},
        )

    entry.refresh_from_db()
    assert generate.call_count == 1
    assert entry.document_title == "The Correct Document Title"
    assert entry.source_filename == "renamed.docx"
    assert "aucun appel API" in response.content.decode()


def test_per_user_minute_limit_blocks_additional_api_call(client, monkeypatch) -> None:
    monkeypatch.setenv("SUMMARY_USER_MINUTE_LIMIT", "1")
    user = _user("limited")
    client.force_login(user)

    with patch(
        "catalog.services.publication_documents.generate_bilingual_summary",
        return_value=_summary("first"),
    ) as generate:
        client.post(
            reverse("document-summary"),
            {"document": _upload(_docx_bytes("first document"))},
        )
        response = client.post(
            reverse("document-summary"),
            {"document": _upload(_docx_bytes("second document"))},
        )

    assert generate.call_count == 1
    assert "Attendez une minute" in response.content.decode()


def test_global_daily_limit_blocks_api_call(client, monkeypatch) -> None:
    monkeypatch.setenv("SUMMARY_GLOBAL_DAILY_LIMIT", "1")
    existing_user = _user("existing")
    current_user = _user("current")
    DocumentSummaryGenerationAttempt.objects.create(
        owner=existing_user,
        document_sha256="a" * 64,
        model_name="test-model",
    )
    client.force_login(current_user)

    with patch(
        "catalog.services.publication_documents.generate_bilingual_summary"
    ) as generate:
        response = client.post(
            reverse("document-summary"),
            {"document": _upload(_docx_bytes("global limit"))},
        )

    generate.assert_not_called()
    assert "limite quotidienne de l’application" in response.content.decode()


def test_active_generation_blocks_concurrent_request(client) -> None:
    user = _user("concurrent")
    ActiveDocumentSummaryGeneration.objects.create(owner=user)
    client.force_login(user)

    with patch(
        "catalog.services.publication_documents.generate_bilingual_summary"
    ) as generate:
        response = client.post(
            reverse("document-summary"),
            {"document": _upload(_docx_bytes("concurrent"))},
        )

    generate.assert_not_called()
    assert "déjà en cours" in response.content.decode()


def test_provider_error_is_sanitized(monkeypatch) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "test-secret-key")
    raw_error = b'{"error":{"message":"internal detail must not escape"}}'
    error = HTTPError(
        "https://api.openai.com/v1/responses",
        429,
        "quota",
        {},
        BytesIO(raw_error),
    )

    with (
        patch("catalog.services.document_summaries.urlopen", side_effect=error),
        patch("catalog.services.document_summaries.time.sleep"),
    ):
        with pytest.raises(DocumentSummaryError) as caught:
            generate_bilingual_summary("source text")

    assert "quota" in str(caught.value)
    assert "internal detail" not in str(caught.value)
    assert "test-secret-key" not in str(caught.value)


def _openai_payload() -> BytesIO:
    result = {
        "abstract_en": "English abstract.",
        "abstract_fr": "Résumé français.",
        "keywords_en": ["memory"],
        "keywords_fr": ["mémoire"],
        "suggested_title": "Memory and Archives",
        "suggested_authors": ["Florence Fix"],
        "suggested_publication_year": 2024,
        "suggested_publication_type": "article",
        "suggested_doi": "",
    }
    return BytesIO(json.dumps({"output_text": json.dumps(result)}).encode())


def test_transient_provider_failure_is_retried(monkeypatch) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    error = HTTPError(
        "https://api.openai.com/v1/responses", 503, "unavailable", {}, BytesIO()
    )

    with (
        patch(
            "catalog.services.document_summaries.urlopen",
            side_effect=[error, _openai_payload()],
        ) as call,
        patch("catalog.services.document_summaries.time.sleep") as sleep,
    ):
        result = generate_bilingual_summary("source text")

    assert result.abstract_en == "English abstract."
    assert call.call_count == 2
    sleep.assert_called_once_with(1.0)


def test_retry_after_header_controls_transient_delay(monkeypatch) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    error = HTTPError(
        "https://api.openai.com/v1/responses",
        429,
        "limited",
        {"Retry-After": "2"},
        BytesIO(),
    )

    with (
        patch(
            "catalog.services.document_summaries.urlopen",
            side_effect=[error, _openai_payload()],
        ),
        patch("catalog.services.document_summaries.time.sleep") as sleep,
    ):
        generate_bilingual_summary("source text")

    sleep.assert_called_once_with(2.0)


def test_network_failure_uses_two_retries_then_fails(monkeypatch) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")

    with (
        patch(
            "catalog.services.document_summaries.urlopen",
            side_effect=URLError("temporary DNS failure"),
        ) as call,
        patch("catalog.services.document_summaries.time.sleep") as sleep,
    ):
        with pytest.raises(DocumentSummaryError, match="momentanément inaccessible"):
            generate_bilingual_summary("source text")

    assert call.call_count == 3
    assert [item.args[0] for item in sleep.call_args_list] == [1.0, 3.0]


def test_non_transient_client_error_is_not_retried(monkeypatch) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    error = HTTPError("https://api.openai.com/v1/responses", 400, "bad", {}, BytesIO())

    with (
        patch("catalog.services.document_summaries.urlopen", side_effect=error) as call,
        patch("catalog.services.document_summaries.time.sleep") as sleep,
    ):
        with pytest.raises(DocumentSummaryError, match="n’a pas accepté"):
            generate_bilingual_summary("source text")

    assert call.call_count == 1
    sleep.assert_not_called()


@pytest.mark.parametrize(
    ("name", "contents", "message"),
    [
        ("fake.pdf", b"this is not a pdf" * 30, "ne correspond pas à un document PDF"),
        ("fake.docx", b"PK but not a zip archive" * 30, "endommagé"),
    ],
)
def test_spoofed_or_malformed_documents_are_rejected(name, contents, message) -> None:
    with pytest.raises(DocumentSummaryError, match=message):
        extract_document_text(_upload(contents, name))


def test_docx_zip_bomb_ratio_is_rejected() -> None:
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "x")
        archive.writestr("word/document.xml", "a" * 2_000_000)

    with pytest.raises(DocumentSummaryError, match="compression"):
        extract_document_text(_upload(output.getvalue()))


def test_encrypted_pdf_is_rejected() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("password")
    output = BytesIO()
    writer.write(output)

    with pytest.raises(DocumentSummaryError, match="mot de passe"):
        extract_document_text(_upload(output.getvalue(), "protected.pdf"))


def test_pdf_page_limit_is_enforced() -> None:
    writer = PdfWriter()
    for _ in range(301):
        writer.add_blank_page(width=100, height=100)
    output = BytesIO()
    writer.write(output)

    with pytest.raises(DocumentSummaryError, match="300 pages"):
        extract_document_text(_upload(output.getvalue(), "long.pdf"))


def test_admin_usage_dashboard_shows_aggregate_user_stats(client) -> None:
    admin_user = get_user_model().objects.create_superuser(
        username="admin-stats",
        password="safe-test-password",
    )
    author = _user("summary-author")
    DocumentSummaryGenerationAttempt.objects.create(
        owner=author,
        document_sha256="b" * 64,
        model_name="test-model",
    )
    DocumentSummaryCache.objects.create(
        owner=author,
        document_sha256="b" * 64,
        model_name="test-model",
        generator_version="test-v1",
        abstract_en="English",
        abstract_fr="Français",
        keywords_en=["history"],
        keywords_fr=["histoire"],
    )
    client.force_login(admin_user)

    response = client.get(
        reverse("admin:catalog_documentsummarygenerationattempt_changelist")
    )

    content = response.content.decode()
    assert response.status_code == 200
    assert "Utilisation des résumés par utilisateur" in content
    assert "summary-author" in content
    assert "Résultats en cache" in content


def _cache_entry(owner, marker: str = "c") -> DocumentSummaryCache:
    return DocumentSummaryCache.objects.create(
        owner=owner,
        document_sha256=marker * 64,
        model_name="test-model",
        generator_version="test-v1",
        abstract_en="English",
        abstract_fr="Français",
        keywords_en=["history"],
        keywords_fr=["histoire"],
    )


def test_owner_can_delete_cache_entry_with_audit_record(client) -> None:
    owner = _user("cache-owner")
    entry = _cache_entry(owner)
    client.force_login(owner)

    response = client.post(
        reverse("document-summary-cache-delete", args=[entry.id])
    )

    assert response.status_code == 302
    assert not DocumentSummaryCache.objects.filter(id=entry.id).exists()
    event = AuditEvent.objects.get(action="document_summary_cache.deleted")
    assert event.actor == owner
    assert event.object_id == str(entry.id)
    assert event.before_checksum == "c" * 64
    assert event.metadata["reason"] == "owner_request"
    assert "abstract" not in str(event.metadata).lower()


def test_owner_can_view_cached_result_but_other_user_cannot(client) -> None:
    owner = _user("detail-owner")
    other = _user("detail-other")
    entry = _cache_entry(owner, "9")
    DocumentSummaryCache.objects.filter(id=entry.id).update(
        source_filename="history-of-memory.docx",
        document_title="A History of Memory",
    )

    client.force_login(owner)
    owner_response = client.get(
        reverse("document-summary-cache-detail", args=[entry.id])
    )
    client.force_login(other)
    other_response = client.get(
        reverse("document-summary-cache-detail", args=[entry.id])
    )

    owner_content = owner_response.content.decode()
    assert owner_response.status_code == 200
    assert "A History of Memory" in owner_content
    assert "history-of-memory.docx" in owner_content
    assert "English" in owner_content
    assert "Français" in owner_content
    assert other_response.status_code == 404


def test_cache_history_shows_friendly_metadata_and_detail_link(client) -> None:
    owner = _user("history-owner")
    entry = _cache_entry(owner, "8")
    DocumentSummaryCache.objects.filter(id=entry.id).update(
        source_filename="friendly.docx",
        document_title="Friendly Document Title",
    )
    client.force_login(owner)

    response = client.get(reverse("document-summary"))

    content = response.content.decode()
    assert response.status_code == 200
    assert "Friendly Document Title" in content
    assert "friendly.docx" in content
    assert reverse("document-summary-cache-detail", args=[entry.id]) in content


def test_user_cannot_delete_another_users_cache_entry(client) -> None:
    owner = _user("actual-owner")
    intruder = _user("other-user")
    entry = _cache_entry(owner)
    client.force_login(intruder)

    response = client.post(
        reverse("document-summary-cache-delete", args=[entry.id])
    )

    assert response.status_code == 404
    assert DocumentSummaryCache.objects.filter(id=entry.id).exists()
    assert not AuditEvent.objects.filter(action="document_summary_cache.deleted").exists()


def test_retention_purges_only_expired_cache_with_audit(monkeypatch) -> None:
    monkeypatch.setenv("SUMMARY_CACHE_RETENTION_DAYS", "90")
    owner = _user("retention-owner")
    expired = _cache_entry(owner, "d")
    current = _cache_entry(owner, "e")
    DocumentSummaryCache.objects.filter(id=expired.id).update(
        created_at=timezone.now() - timedelta(days=91)
    )

    count = purge_expired_summary_cache()

    assert count == 1
    assert not DocumentSummaryCache.objects.filter(id=expired.id).exists()
    assert DocumentSummaryCache.objects.filter(id=current.id).exists()
    event = AuditEvent.objects.get(action="document_summary_cache.deleted")
    assert event.actor is None
    assert event.metadata["reason"] == "retention"


def test_retention_command_is_dry_run_unless_apply_is_passed(capsys) -> None:
    owner = _user("command-owner")
    expired = _cache_entry(owner, "f")
    DocumentSummaryCache.objects.filter(id=expired.id).update(
        created_at=timezone.now() - timedelta(days=91)
    )

    call_command("purge_summary_cache")
    assert DocumentSummaryCache.objects.filter(id=expired.id).exists()

    call_command("purge_summary_cache", "--apply")
    assert not DocumentSummaryCache.objects.filter(id=expired.id).exists()
    assert "would delete 1" in capsys.readouterr().out
